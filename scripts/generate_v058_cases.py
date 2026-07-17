from __future__ import annotations

import argparse
import json
import shutil
import sys
import zipfile
from dataclasses import asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, PngImagePlugin

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.sample_v058_provisions import sample_provisions
from scripts.v058_case_catalog import (
    SYNTHETIC_BANNER,
    CaseSpec,
    MaterialSpec,
    build_case_catalog,
)


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DATABASE = ROOT / "legal_knowledge" / "index" / "legal_kb.sqlite3"
DEFAULT_OUTPUT = ROOT / "测试用例" / "v058随机条文"
FIXED_DOCUMENT_TIME = datetime(2026, 7, 14, tzinfo=timezone.utc)


def generate_corpus(
    output_root: str | Path,
    *,
    overwrite: bool = False,
    database: str | Path = DEFAULT_DATABASE,
) -> list[Path]:
    destination = Path(output_root).resolve()
    destination.mkdir(parents=True, exist_ok=True)
    manifest = sample_provisions(database, seed=58)
    _write_json(destination / "sampling_manifest.json", manifest)
    provisions = _provision_lookup(manifest)

    generated: list[Path] = []
    for case in build_case_catalog():
        case_dir = destination / f"{case.case_id}_{case.short_name}"
        _prepare_case_directory(case_dir, destination, overwrite=overwrite)
        provision = provisions[(case.law_key, case.article)]
        _write_case_package(case_dir, case, provision, manifest)
        generated.append(case_dir)
    return generated


def _prepare_case_directory(case_dir: Path, output_root: Path, *, overwrite: bool) -> None:
    resolved_case = case_dir.resolve()
    if resolved_case.parent != output_root.resolve():
        raise ValueError(f"Refusing to generate outside output root: {resolved_case}")
    if case_dir.exists():
        if not overwrite:
            raise FileExistsError(f"Case directory already exists: {case_dir}")
        shutil.rmtree(case_dir)
    for folder in ("statements", "reports", "report_images", "expected"):
        (case_dir / folder).mkdir(parents=True, exist_ok=True)


def _provision_lookup(manifest: dict[str, object]) -> dict[tuple[str, str], dict[str, Any]]:
    result: dict[tuple[str, str], dict[str, Any]] = {}
    accepted = manifest["accepted"]
    assert isinstance(accepted, dict)
    for law_key, items in accepted.items():
        assert isinstance(items, list)
        for item in items:
            assert isinstance(item, dict)
            result[(str(law_key), str(item["article"]))] = item
    return result


def _write_case_package(
    case_dir: Path,
    case: CaseSpec,
    provision: dict[str, Any],
    manifest: dict[str, object],
) -> None:
    material_records: list[dict[str, str]] = []
    assertion_records: list[dict[str, object]] = []

    for material in case.materials:
        relative_path = Path(material.folder) / material.filename
        document_path = case_dir / relative_path
        if material.folder == "statements":
            _write_statement_docx(document_path, case, material)
        else:
            _write_report_docx(document_path, case, material)
            image_path = case_dir / "report_images" / f"{document_path.stem}_page_001.png"
            _render_report_png(image_path, case, material)

        material_records.append(
            {
                "material_id": material.material_id,
                "material_type": material.material_type,
                "path": relative_path.as_posix(),
                "title": material.title,
                "role": material.role,
            }
        )
        for index, assertion in enumerate(material.assertions, start=1):
            agent_fact = dict(assertion)
            agent_fact["fact_id"] = f"F-{material.material_id}-{index:02d}"
            agent_fact["person"] = str(
                agent_fact.get("declarant") or agent_fact.get("actor") or ""
            )
            assertion_records.append(
                {
                    "assertion_id": f"A-{material.material_id}-{index:02d}",
                    "source_material_id": material.material_id,
                    "source_path": relative_path.as_posix(),
                    "agent_fact": agent_fact,
                }
            )

    case_payload = {
        "schema_version": "0.58",
        "case_id": case.case_id,
        "short_name": case.short_name,
        "synthetic_notice": SYNTHETIC_BANNER,
        "law_key": case.law_key,
        "law_title": case.law_title,
        "article": case.article,
        "provision_text": provision["text"],
        "complexity_type": case.complexity_type,
        "case_type": case.case_type_hint,
        "summary": case.summary,
        "materials": material_records,
        "authority_verifications": list(case.authority_verifications),
    }
    _write_json(case_dir / "case.json", case_payload)
    _write_json(
        case_dir / "sampling.json",
        {
            "schema_version": "1.0",
            "synthetic_notice": SYNTHETIC_BANNER,
            "seed": manifest["seed"],
            "pool_size": manifest["pool_sizes"][case.law_key],
            "law_key": case.law_key,
            "selected_provision": provision,
        },
    )
    _write_json(
        case_dir / "expected" / "semantic_assertions.json",
        {
            "schema_version": "0.58",
            "case_id": case.case_id,
            "synthetic_notice": SYNTHETIC_BANNER,
            "assertions": assertion_records,
        },
    )
    _write_json(
        case_dir / "expected" / "expected_outcome.json",
        {
            "schema_version": "0.58",
            "case_id": case.case_id,
            "synthetic_notice": SYNTHETIC_BANNER,
            "required_claims": list(case.required_claims),
            "forbidden_claims": list(case.forbidden_claims),
            "expected_bayesian_models": list(case.expected_bayesian_models),
            "expected_abstention": case.expected_abstention,
            "required_legal_articles": list(case.required_legal_articles),
            "required_issue_types": list(case.required_issue_types),
            "forbidden_final_conclusions": list(case.forbidden_final_conclusions),
        },
    )


def _new_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for name in ("Title", "Heading 1", "Heading 2"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    document.core_properties.author = "Va1ha11a v0.58 synthetic corpus generator"
    document.core_properties.created = FIXED_DOCUMENT_TIME
    document.core_properties.modified = FIXED_DOCUMENT_TIME
    return document


def _add_banner(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(SYNTHETIC_BANNER)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    run.font.size = Pt(10)


def _write_statement_docx(path: Path, case: CaseSpec, material: MaterialSpec) -> None:
    document = _new_document()
    _add_banner(document)
    title = document.add_heading(material.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"测试案例编号：{case.case_id}")
    document.add_paragraph(f"材料编号：{material.material_id}")
    document.add_paragraph(f"记录对象角色：{material.role}")
    document.add_paragraph("记录说明：以下问答均为合成情境，用于验证语义抽取和证据边界。")
    for question, answer in material.questions:
        q_paragraph = document.add_paragraph()
        q_run = q_paragraph.add_run(f"问：{question}")
        q_run.bold = True
        document.add_paragraph(f"答：{answer}")
    document.add_paragraph("核对：上述合成记录已按测试情境核对。")
    document.save(path)
    _normalize_docx_archive(path)


def _write_report_docx(path: Path, case: CaseSpec, material: MaterialSpec) -> None:
    document = _new_document()
    _add_banner(document)
    title = document.add_heading(material.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"测试案例编号：{case.case_id}")
    document.add_paragraph(f"报告编号：{material.material_id}")
    document.add_heading("一、材料范围", level=1)
    document.add_paragraph("本报告仅汇总本合成测试包内的记录，不代表任何真实机构结论。")
    document.add_heading("二、核查记录", level=1)
    for index, paragraph in enumerate(material.paragraphs, start=1):
        document.add_paragraph(f"{index}. {paragraph}")
    document.add_heading("三、结论边界", level=1)
    document.add_paragraph(
        "本报告只说明材料中可观察、可核对的事实及其局限，不认定违法犯罪成立，"
        "不决定法律适用或者处罚。"
    )
    document.add_paragraph("测试出具单位：测试记录机构A（测试专用）")
    document.save(path)
    _normalize_docx_archive(path)


def _normalize_docx_archive(path: Path) -> None:
    normalized = path.with_suffix(".normalized.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        normalized,
        "w",
        compression=zipfile.ZIP_DEFLATED,
        compresslevel=9,
    ) as destination:
        for source_info in sorted(source.infolist(), key=lambda item: item.filename):
            target_info = zipfile.ZipInfo(
                source_info.filename,
                date_time=(2026, 7, 14, 0, 0, 0),
            )
            target_info.compress_type = zipfile.ZIP_DEFLATED
            target_info.external_attr = source_info.external_attr
            target_info.create_system = source_info.create_system
            destination.writestr(target_info, source.read(source_info.filename))
    normalized.replace(path)


def _render_report_png(path: Path, case: CaseSpec, material: MaterialSpec) -> None:
    width, height = 1240, 1754
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = _find_cjk_font()
    title_font = ImageFont.truetype(str(font_path), 46)
    heading_font = ImageFont.truetype(str(font_path), 31)
    body_font = ImageFont.truetype(str(font_path), 26)
    banner_font = ImageFont.truetype(str(font_path), 24)

    y = 60
    y = _draw_centered(draw, SYNTHETIC_BANNER, y, banner_font, fill="#b00020", width=width)
    y += 45
    y = _draw_centered(draw, material.title, y, title_font, fill="black", width=width)
    y += 44
    draw.text((100, y), f"测试案例编号：{case.case_id}", font=body_font, fill="black")
    y += 48
    draw.text((100, y), f"报告编号：{material.material_id}", font=body_font, fill="black")
    y += 70
    draw.line((90, y, width - 90, y), fill="#333333", width=2)
    y += 40
    draw.text((100, y), "核查记录", font=heading_font, fill="black")
    y += 58
    for index, paragraph in enumerate(material.paragraphs, start=1):
        for line in _wrap_cjk(f"{index}. {paragraph}", 35):
            draw.text((110, y), line, font=body_font, fill="#202020")
            y += 43
        y += 16
    y += 20
    draw.text((100, y), "结论边界", font=heading_font, fill="black")
    y += 58
    boundary = "仅描述合成材料中的事实和局限，不认定违法犯罪成立，不决定法律适用或处罚。"
    for line in _wrap_cjk(boundary, 35):
        draw.text((110, y), line, font=body_font, fill="#202020")
        y += 43

    draw.ellipse((width - 330, height - 330, width - 100, height - 100), outline="#bb2233", width=5)
    stamp = "测试专用"
    stamp_box = draw.textbbox((0, 0), stamp, font=heading_font)
    stamp_width = stamp_box[2] - stamp_box[0]
    draw.text((width - 215 - stamp_width / 2, height - 230), stamp, font=heading_font, fill="#bb2233")
    draw.rectangle((25, 25, width - 25, height - 25), outline="#555555", width=2)

    metadata = PngImagePlugin.PngInfo()
    metadata.add_text("Synthetic-Notice", SYNTHETIC_BANNER)
    metadata.add_text("Case-ID", case.case_id)
    image.save(path, pnginfo=metadata, optimize=False)


def _draw_centered(
    draw: ImageDraw.ImageDraw,
    text: str,
    y: int,
    font: ImageFont.FreeTypeFont,
    *,
    fill: str,
    width: int,
) -> int:
    box = draw.textbbox((0, 0), text, font=font)
    text_width = box[2] - box[0]
    draw.text(((width - text_width) / 2, y), text, font=font, fill=fill)
    return y + box[3] - box[1]


def _wrap_cjk(text: str, width: int) -> list[str]:
    return [text[index : index + width] for index in range(0, len(text), width)] or [""]


def _find_cjk_font() -> Path:
    candidates = (
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
    )
    for path in candidates:
        if path.is_file():
            return path
    raise FileNotFoundError("No supported CJK font found (Microsoft YaHei, SimSun, or SimHei).")


def _write_json(path: Path, value: object) -> None:
    path.write_text(
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate the v0.58 synthetic ten-case corpus.")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--database", type=Path, default=DEFAULT_DATABASE)
    parser.add_argument("--overwrite", action="store_true")
    args = parser.parse_args(argv)
    generated = generate_corpus(
        args.output,
        overwrite=args.overwrite,
        database=args.database,
    )
    print(json.dumps({"output": str(args.output.resolve()), "case_count": len(generated)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
